# -*- coding: utf-8 -*-
"""
生成《实验设置说明》Word 文档
============================
把实验环境 / 实验参数 (含关键训练超参) / 实验数据 / 对比基线 /
HV 计算口径 / 三张图配文 写成一份 .docx, 并内嵌三张图.
输出: 实验设置说明.docx
"""
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.abspath(__file__))
PIC = os.path.join(ROOT, 'pic')
OUT = os.path.join(ROOT, '实验设置说明.docx')

CJK = '宋体'
LATIN = 'Times New Roman'


def set_cjk(run):
    run.font.name = LATIN
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), CJK)


def body(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_cjk(run)
    return p


def heading(doc, text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    set_cjk(run)
    return p


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    set_cjk(run)


def add_image(doc, fname, width_in):
    path = os.path.join(PIC, fname)
    if os.path.isfile(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_in))


def main():
    doc = Document()
    # 默认样式字体
    st = doc.styles['Normal']
    st.font.name = LATIN
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)

    title(doc, '实验设置说明')

    heading(doc, '一、实验环境')
    body(doc,
         '本文实验在一台配备 Intel Core i5-12500H 处理器与 NVIDIA RTX 3060 GPU '
         '的 Windows 11 工作站上完成,深度学习框架采用 PyTorch 2.5.1,Python 版本为 '
         '3.9.12,所有算法的策略网络训练与推理均在 GPU 上执行,而仿真环境与经验回放'
         '缓冲区维护在 CPU 端。为保证多方法对比的公平与可复现性,实验对底层数学库做了'
         '单线程化设置以消除并行非确定性对超体积计算的干扰,并通过统一的评测种子机制,'
         '让所有参与对比的方法在完全相同的任务到达序列、信道衰落实现与服务器配置上接受'
         '评测,从而确保不同算法之间的性能差异只来源于策略本身,而非随机环境实现的差异。')

    heading(doc, '二、实验参数')
    body(doc,
         '仿真环境模拟边缘-云端协同计算场景,系统由 6 个异构边缘服务器与 1 个云端节点'
         '构成,以 1 秒为一个系统时隙离散推进,并在动作空间上将云端与边缘统一编码——'
         '动作 0 指向云端、动作 1 至 6 分别指向各边缘服务器,以对齐 GMORL 协议下的基线'
         '对比口径。在算力配置上,边缘服务器的 CPU 频率在每个回合从 10 至 40 GHz 区间'
         '均匀采样,云端节点频率更高、取自 50 至 70 GHz,体现数据中心的高主频特性;'
         '任务的数据量在 10 至 40 Mbit 范围内均匀采样,每个任务的计算密度则固定在 '
         '0.1024 至 0.3072 Gigacycles/Mbit 之间。能耗方面,边缘服务器的有效电容系数取 '
         '1×10⁻³,云端取 1×10⁻⁴,即为边缘的十分之一,这一更低的系数用以反映数据中心'
         '定制硬件、液冷优化以及静态功率被多租户分摊后的等效效果,使云端在更高频率下依然'
         '保持能耗竞争力。信道采用瑞利衰落模型,每个时隙、任务与服务器组合独立采样小尺度'
         '衰落增益并裁剪至 0.3 至 2.5 之间,上行可达速率由香农公式计算且其瞬时信噪比受该'
         '衰落增益调制,经带宽与噪声标定后,边缘链路的有效传输速率落在 400 至 500 Mbps、'
         '云端链路因物理距离更远而落在 80 至 120 Mbps;卸载发射功率设为 0.5 W。在此设置'
         '下,任务的端到端时延由传输、计算与排队三部分构成,能耗由传输能耗与正比于频率'
         '平方的执行能耗构成,云端的低电容系数与边缘的低传输代价共同构成"云端算力强但'
         '传输慢、边缘传输快但算力弱"的真实多目标冲突。偏好测试集在二维单纯形上均匀采样,'
         '共 21 个偏好点,从极端重能耗(延迟权重为 0)以 0.05 为步长递增至极端重延迟'
         '(延迟权重为 1),完整覆盖各种偏好组合。')
    body(doc,
         '在算法训练设置上,PC-FDN 以带向量化 Q 评估与冲突目标正则的 Soft Actor-Critic '
         '为主干,并以偏好超网络替代查表式先验、以 Normalize-and-Project 机制对回报量级'
         '做在线归一,所有网络的隐藏层宽度均为 128。优化器统一采用 Adam,其中策略网络的'
         '学习率为 1×10⁻⁴、双 Q 评论家网络的学习率为 1×10⁻³、自动温度系数的学习率为 '
         '3×10⁻⁴;折扣因子取 0.95,目标网络以系数 0.005 做软更新。熵正则温度采用自动'
         '调节机制,初始值设为 0.05、目标熵设为 0.5(约为离散动作最大熵 log6 的三成),'
         '以避免在离散动作空间下温度退化为零而导致策略坍缩。经验回放缓冲区容量为 1 万条,'
         '批大小为 64,在缓冲区累积满 500 条样本后开始参数更新且每 4 个环境步更新一次,'
         '扩散式反馈策略的去噪步数设为 3。此外,冲突目标正则与策略多样性正则的系数均取 '
         '0.1,偏好先验缓冲区的指数滑动平均衰减系数为 0.5、检索扰动幅度为 0.05。在训练'
         '规模上,每个方法训练 100 个 epoch、每个 epoch 在 8 个随机偏好上采样,单回合的'
         '时序长度为 100 个时隙;最终评测在 21 个均匀偏好上各重复 3 个回合并取聚合结果。')

    heading(doc, '三、实验数据与任务到达')
    body(doc,
         '实验中节点的异构物理参数与任务到达模式以阿里巴巴 OpenB 生产集群调度追踪数据集'
         '为依据进行标定,该数据集涵盖数千异构节点的资源配置以及数十万容器任务的全生命'
         '周期调度记录,为仿真环境提供了贴近真实生产负载的分布依据。在此基础上,任务到达'
         '过程建模为非平稳泊松流,即每个时隙到达的任务数服从泊松分布,而其到达率随时间'
         '变化,在一个基础到达率之上叠加一个正弦周期调制项,用以刻画真实数据中心中任务'
         '负载随昼夜节律起伏的周期性波动,从而使仿真不仅在静态分布上、也在时间动态上逼近'
         '真实边缘计算场景。')

    heading(doc, '四、对比基线')
    body(doc,
         '为充分检验所提方法的有效性,本文与一组覆盖强化学习、上下文 bandit、进化算法与'
         '启发式策略的基线进行对比,其中除进化与退火类方法外,所有学习型基线均与 PC-FDN '
         '采用同一"单策略且偏好作为上下文输入"的训练范式,在 21 个偏好上联合训练并在同'
         '一偏好集上评测,以保证比较的同源性。具体而言,PC-FDN 为本文所提方法,其核心是'
         '偏好条件化的策略生成与由缓冲区驱动的测试期自适应;GenMOSAC 对齐可泛化卸载工作,'
         '采用一个基于 Deep Sets 思想的置换不变状态编码器,即对每个边缘服务器施加共享的'
         '多层感知机后做带掩码的最大池化聚合,并辅以队列负载直方图与任务/偏好两路编码拼接,'
         '其训练主干为带双 Q 网络和自动温度调节的离散 Soft Actor-Critic,需要强调的是该'
         '方法并不使用任何注意力机制或 Set-Transformer,它与下面的 Discrete SAC 基线的'
         '唯一区别就在于状态编码器;Discrete SAC 直接将拼平的状态向量送入多层感知机,以'
         '偏好对奖励做线性标量化;LDQN 在深度 Q 网络中引入 LSTM 单元,在回合内保持隐状态'
         '以捕捉任务到达的时序规律;LinUCB 是线性上下文 bandit,将任务、偏好与所选服务器'
         '的特征拼成上下文向量并按置信上界探索,但不建模后续时序回报;NSGA-II 是经典的'
         '进化多目标算法,以时延与能耗为双目标通过非支配排序与拥挤度距离搜索 Pareto 解集;'
         'SA 为模拟退火,以偏好加权的标量化代价按 Metropolis 准则随温度退火搜索,并对每个'
         '偏好独立求解;此外还包含在合法动作上均匀随机选择的 Random 策略作为性能下界,以及'
         '在仅边缘环境中按轮询顺序分配的 Round-Robin 策略作为无学习的朴素参照。')

    heading(doc, '五、Pareto 前沿与超体积计算')
    body(doc,
         '性能评测以 Pareto 前沿与超体积(Hypervolume, HV)指标为核心。对每个方法在 21 '
         '个偏好下分别评测,聚合得到 21 个以平均时延与平均能耗为坐标的二维点,再对其做'
         '二维最小化意义下的非支配排序,即沿时延升序扫描并仅保留能耗严格下降的点,得到'
         '该方法的 Pareto 前沿。为使不同方法的 HV 严格可比,本文将所有方法的全部评测点'
         '汇集后取各维度的最大值并放大 1.05 倍,构造一个全体方法共用的全局参考点,再以此'
         '参考点计算每条 Pareto 前沿与参考点围成的覆盖面积作为 HV。HV 越大,意味着该方法'
         '在统一参考下越靠近理想的低时延、低能耗角点、所覆盖的目标空间面积越大,综合多目'
         '标性能越强。')

    heading(doc, '六、图表配文')
    add_image(doc, 'fig_omega_mofd.png', 5.2)
    body(doc,
         '图 1 展示了所提 PC-FDN 在边缘-云协同环境下对偏好信号的响应特性:横轴为延迟偏好'
         '权重,从极端重能耗一侧逐步扫向极端重延迟一侧,左纵轴对应平均任务完成时延、右纵'
         '轴对应平均能耗,两条曲线随偏好权重单调地反向变化,平均时延由约 110 秒下降至约 '
         '28 秒,平均能耗则由约 2.6 焦耳上升至约 4.0 焦耳,说明该方法能够将连续变化的偏好'
         '信号忠实地解码为连续可控的时延-能耗权衡,而不会坍缩到单一固定的工作点。')

    add_image(doc, 'pareto_comparison.png', 5.0)
    body(doc,
         '图 2 给出了在仅边缘环境下、统一参考点下各方法的 Pareto 前沿对比,横纵轴分别为'
         '平均时延与平均能耗,所提 PC-FDN 与 GenMOSAC、Discrete SAC、LDQN、LinUCB、SA、'
         'NSGA-II、Random 和 Round-Robin 共八个基线在 21 个均匀偏好上同源评测,各方法的'
         '评测点经非支配排序后形成 Pareto 前沿,并在全体方法共用的全局参考点下计算 HV;'
         '结果显示 PC-FDN 的前沿在整个偏好区间持续位于左下方,同时把时延与能耗推向更优的'
         '角点,而 Random、Round-Robin、LinUCB 等启发式与 bandit 方法则集中在右上方的被'
         '支配区域,差距明显。')

    add_image(doc, 'v2_pareto_comparison.png', 6.2)
    body(doc,
         '图 3 针对引入云端节点的边缘-云协同环境给出 Pareto 前沿对比,此设置在动作空间中'
         '加入云端选项,从而在云端与边缘之间形成"快但耗能、慢但省能"的真实冲突以对齐 '
         'GMORL 协议;参与对比的方法与图 2 基本一致,均以偏好作为上下文输入并在同源评测下'
         '得到各自的 Pareto 前沿。图中左侧子图为完整视图,可见 PC-FDN 勾勒出最内层的包络'
         '并显著优于全部基线;右侧子图则对低时延的稠密区域做局部放大,用以分辨膝点附近各'
         '方法的相对排序,进一步表明 PC-FDN 的优势贯穿整个偏好谱而非仅出现在极端偏好处;'
         '两个子图共用同一参考点,因此其超体积差距可以直接横向比较。')

    add_image(doc, 'hv_comparison_bar.png', 6.4)
    body(doc,
         '图 4 以柱状图的形式直接对比了两种环境下各方法的超体积:左面板为仅边缘环境、'
         '右面板为边缘-云环境,每个面板内按 HV 由高到低排列,PC-FDN(ours)以红色高亮、'
         '其余基线统一为蓝色,柱顶标注各方法的 HV 数值。可以看到 PC-FDN 在两种环境下均取得'
         '最高的超体积,验证了其在不同协同计算场景下的一致优势;需要说明的是两个面板的纵轴'
         '量级不同,源于两套环境的目标量级与参考点本身存在差异,因此 HV 数值只应在各自环境'
         '内横向比较,而不宜跨环境直接比较绝对值。')

    doc.save(OUT)
    print(f'[save] {OUT}')


if __name__ == '__main__':
    main()
